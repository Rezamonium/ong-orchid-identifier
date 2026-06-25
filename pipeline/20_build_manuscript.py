"""
20_build_manuscript.py — Assemble the ViT-vs-CNN orchid manuscript as a styled .docx.

Style: Times New Roman 11 pt, justified body, 1.5 line spacing, Word Heading 1/2/3
(black) so the Navigation pane is populated, embedded figures + tables, author-year
citations and an alphabetical reference list, page numbers in the footer.

Numbers are taken from the clean-pilot eval outputs (colab/result/eval/...). Figures are
embedded from colab/result/eval/_comparison/figures/.

Run:  python notebooks/20_build_manuscript.py
Out:  article/ONG_orchid_ViT_vs_CNN_manuscript.docx
"""

import sys
from pathlib import Path
try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ONG = Path(r"E:/Claude Code/ONG_v3")
FIG = ONG / "colab/result/eval/_comparison/figures"
OUT = ONG / "article/ONG_orchid_ViT_vs_CNN_manuscript.docx"
FONT = "Times New Roman"

doc = Document()

# ── base styles ───────────────────────────────────────────────────────────────────
def _set_font(style, size, bold=None, italic=None, color=None):
    style.font.name = FONT
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic
    if color is not None:
        style.font.color.rgb = color
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), FONT)
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)

normal = doc.styles["Normal"]
_set_font(normal, 11)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after = Pt(6)

BLACK = RGBColor(0, 0, 0)
for name, sz in [("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)]:
    h = doc.styles[name]
    _set_font(h, sz, bold=True, color=BLACK)
    h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
doc.styles["Heading 3"].font.italic = True


# ── helpers ───────────────────────────────────────────────────────────────────────
def para(text="", *, align="justify", bold=False, italic=False, size=11, after=6, before=0):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = {
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY, "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if text:
        r = p.add_run(text)
        r.bold, r.italic, r.font.size = bold, italic, Pt(size)
    return p


def body(text):
    """Paragraph that supports **bold** and *italic* inline markup."""
    import re
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for tok in re.split(r"(\*\*.+?\*\*|\*.+?\*)", text):
        if not tok:
            continue
        r = p.add_run(tok[2:-2] if tok.startswith("**") else tok[1:-1] if tok.startswith("*") else tok)
        r.bold = tok.startswith("**")
        r.italic = tok.startswith("*") and not tok.startswith("**")
    return p


def figure(fname, caption, width=6.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.add_run().add_picture(str(FIG / fname), width=Inches(width))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    c.paragraph_format.space_after = Pt(10)
    runs = caption.split("**", 2)
    c.add_run(runs[0]).bold = True            # "Figure N."
    if len(runs) >= 3:
        c.add_run(runs[1]).bold = True
        tail = c.add_run(runs[2])
    else:
        tail = c.add_run("")
    for r in c.runs:
        r.font.size = Pt(10)


def table(headers, rows, caption, col0_left=True, fontsize=9):
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(8)
    rc = cap.add_run(caption.split(" ", 2)[0] + " " + caption.split(" ", 2)[1] + " ")
    rc.bold = True
    cap.add_run(caption.split(" ", 2)[2]).bold = False
    for r in cap.runs:
        r.font.size = Pt(10)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(fontsize)
        run.font.name = FONT
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ""
            pr = cells[j].paragraphs[0]
            run = pr.add_run(str(val))
            run.font.size = Pt(fontsize)
            run.font.name = FONT
            pr.alignment = WD_ALIGN_PARAGRAPH.LEFT if (j == 0 and col0_left) else WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_page_numbers():
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for t, attr in [("begin", "w:fldCharType"), (None, None), ("end", "w:fldCharType")]:
        if t in ("begin", "end"):
            fld = OxmlElement("w:fldChar"); fld.set(qn(attr), t); run._r.append(fld)
        else:
            it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = "PAGE"
            run._r.append(it)
    run.font.name = FONT; run.font.size = Pt(10)


# ── FRONT MATTER ──────────────────────────────────────────────────────────────────
para("Vision Transformers versus convolutional networks for fine-grained orchid genus "
     "identification in a species-rich, data-poor flora: a controlled benchmark on the "
     "Orchidaceae of New Guinea", align="center", bold=True, size=15, after=10)

para("Reza Saputra¹,²,*, Kurt Metzger³, Jeffrey Champion⁴, André Schuiteman⁵, "
     "Ashley Field⁶,⁷, Katharina Nargar⁷,⁸, William Edwards²", align="center", size=11, after=4)
para("¹ Southwest Papua Natural Resources Conservation Agency, Ministry of Forestry, Indonesia",
     align="center", size=10, after=0)
para("² James Cook University, Cairns, Queensland, Australia", align="center", size=10, after=0)
para("³ Aiyura, Eastern Highlands Province, Papua New Guinea", align="center", size=10, after=0)
para("⁴ Rumahku di Awan, Candikuning, Baturiti, Tabanan, Bali 82191, Indonesia",
     align="center", size=10, after=0)
para("⁵ Royal Botanic Gardens, Kew, Richmond, Surrey, UK", align="center", size=10, after=0)
para("⁶ Queensland Herbarium, Brisbane, Queensland, Australia", align="center", size=10, after=0)
para("⁷ Australian Tropical Herbarium, Cairns, Queensland, Australia", align="center", size=10, after=0)
para("⁸ CSIRO, Australia", align="center", size=10, after=2)
para("* Correspondence: reza.saputraa21@gmail.com", align="center", size=10, after=10)

doc.add_heading("Abstract", level=1)
body(
    "Identifying orchids from photographs is a paradigmatic challenge for automated "
    "biodiversity informatics: New Guinea, the world's richest island flora, harbours more "
    "than 2,900 orchid species, yet most species are represented by only a handful of "
    "photographs—far fewer than direct species-level classification requires. We address "
    "this with a two-stage system that first predicts the *genus* of a query photograph and "
    "then retrieves visually similar reference images of candidate *species*. Central to the "
    "study is a controlled comparison of four pretrained backbones—two Vision Transformers "
    "(DINOv2, BioCLIP 2) and two modern convolutional networks (ConvNeXt V2-L, "
    "EfficientNetV2-L)—fine-tuned under an identical protocol on a single, frozen, "
    "species-stratified partition of 16,701 live field photographs spanning 120 genera and "
    "1,350 species. On the held-out test set, the self-supervised Vision Transformer DINOv2 "
    "attained the best genus performance (macro-averaged top-1 accuracy 66.9%, 95% CI "
    "63.7–70.6; global top-1 88.9%), and both Vision Transformers outranked both "
    "convolutional networks with non-overlapping confidence intervals. Notably, general-purpose "
    "self-supervised pretraining (DINOv2) outperformed domain-matched biological "
    "vision–language pretraining (BioCLIP 2) by 7.1 points of macro top-1. Errors were "
    "concentrated in two abundant “magnet” genera (*Dendrobium*, *Bulbophyllum*) that "
    "absorbed most misclassifications of low-support genera. For species retrieval, DINOv2 "
    "embeddings achieved species Recall@5 of 86.6% and genus Recall@5 of 98.7%, and their "
    "two-dimensional projection revealed clear genus structure that justifies nearest-neighbour "
    "retrieval. In their as-trained outputs the Vision Transformers were less well calibrated "
    "than the convolutional networks—a systematic under-confidence that post-hoc temperature "
    "scaling removed, lowering every backbone's Expected Calibration Error to ≤0.03 without "
    "changing accuracy. The deployed system is released "
    "as an open web application linked to a multi-access identification key for expert "
    "verification. Our results "
    "indicate that a self-supervised Vision-Transformer backbone combined with embedding "
    "retrieval is an effective, deployable strategy for fine-grained identification in "
    "species-rich, data-poor floras."
)
body("**Keywords:** fine-grained image classification; Vision Transformer; self-supervised "
     "learning; image retrieval; Orchidaceae; New Guinea; biodiversity informatics; data scarcity")

# ── 1. INTRODUCTION ───────────────────────────────────────────────────────────────
doc.add_heading("1. Introduction", level=1)
body(
    "New Guinea is the most floristically diverse island on Earth, with an expert-verified "
    "checklist of 13,634 vascular plant species, 68% of them endemic (Cámara-Leret et al., "
    "2020). The Orchidaceae are its single largest family: more than 2,900 species in roughly "
    "133 genera have been documented (Schuiteman, 1995; de Vogel et al., 2014), making the "
    "region a global hotspot of orchid evolutionary distinctiveness and a priority for "
    "conservation (Vitt et al., 2023). Identifying these orchids is difficult even for "
    "specialists—many genera are separated by subtle floral micro-characters—and expert "
    "capacity is scarce relative to the scale of undescribed and under-documented diversity."
)
body(
    "Automated, image-based identification has transformed botanical practice over the past "
    "decade. Citizen-science platforms such as Pl@ntNet now serve hundreds of millions of "
    "identifications using deep convolutional neural networks (Affouard et al., 2017), and "
    "purpose-built classifiers have been reported for orchids specifically, typically on small "
    "datasets of a few hundred to a few thousand images and a handful of species "
    "(Arwatchananukul et al., 2020; Apriyanti et al., 2023). However, these systems generally "
    "assume that each target class is represented by enough labelled images to train a direct "
    "species classifier. In a flora such as New Guinea's that assumption fails: the average "
    "species is represented by fewer than five photographs, a regime in which per-species "
    "classification is statistically infeasible and the class distribution is severely "
    "long-tailed (Van Horn et al., 2018; Cui et al., 2019)."
)
body(
    "Two further questions remain unsettled for fine-grained botanical recognition. First, "
    "*which backbone architecture* should anchor such a system? Vision Transformers (Dosovitskiy "
    "et al., 2021) now rival or exceed convolutional networks on many benchmarks, and "
    "self-supervised Transformers such as DINOv2 (Oquab et al., 2024) produce general-purpose "
    "features without labels, while modern convolutional networks such as ConvNeXt V2 (Woo et "
    "al., 2023) and EfficientNetV2 (Tan & Le, 2021) remain strong and efficient. Second, does "
    "*domain-matched pretraining*—for example the biological vision–language model BioCLIP "
    "and its successor BioCLIP 2 (Stevens et al., 2024; Gu et al., 2025)—outperform "
    "general-purpose pretraining for a narrow taxonomic task? These questions are rarely "
    "addressed under controlled conditions on the same data."
)
body(
    "Here we present a two-stage identification system for the orchids of New Guinea and use it "
    "as the setting for a controlled architecture benchmark. Our contributions are: (1) a "
    "two-stage pipeline (Figure 1) that combines genus-level classification with FAISS-based species "
    "retrieval (Johnson et al., 2021), designed explicitly for data-scarce floras; (2) a "
    "controlled comparison of two Vision-Transformer and two convolutional backbones, fine-tuned "
    "under an identical protocol on a single frozen, species-stratified partition, that "
    "disentangles the effect of architecture family from that of pretraining domain; (3) a "
    "detailed characterisation of the best model, including per-genus accuracy, calibration, the "
    "long-tailed “magnet-class” error structure, and the geometry of its embedding space; "
    "and (4) an open, deployed web application integrated with a multi-access "
    "identification key for expert verification."
)

# ── 2. METHODS ────────────────────────────────────────────────────────────────────
doc.add_heading("2. Materials and Methods", level=1)
body(
    "Figure 1 gives an overview of the two-stage system. Stage 1 predicts the genus of a query "
    "photograph with a fine-tuned Vision-Transformer backbone; Stage 2 restricts a FAISS "
    "similarity index to the predicted genus and retrieves the most similar reference photographs "
    "as species suggestions. An open-set distance gate between the two stages flags probable "
    "novel genera for expert review rather than forcing a confident misclassification."
)
figure("figure1_ed2.png",
       "**Figure 1.** Overview of the two-stage orchid identification pipeline. In Stage 1 "
       "(genus classification), a DINOv2 ViT-L/14 backbone (fine-tuned, 448 px) feeds a linear "
       "classifier head over the 120 genera followed by softmax with post-hoc temperature scaling, "
       "yielding the predicted genus (top-1) and a calibrated confidence. The backbone's "
       "pre-classifier, L2-normalised 1024-dimensional embedding plays two roles: an open-set gate "
       "scores its cosine distance to the nearest known reference embedding—queries beyond a "
       "threshold are flagged as a probable novel/unknown genus and routed to expert review through "
       "a multi-access key—while the same query embedding drives Stage 2. In Stage 2 (species "
       "retrieval), an exact FAISS inner-product index, restricted to the predicted genus (dashed "
       "arrow), returns the top-k nearest reference photographs and their species labels as ranked "
       "species suggestions for the user's final visual determination.", width=6.8)

doc.add_heading("2.1 Dataset", level=2)
body(
    "We compiled 16,701 live field photographs of New Guinea orchids spanning 120 genera and "
    "1,350 species, drawn from curated public and contributed sources (orchidsnewguinea.com and "
    "associated collections, supplemented by field photography). Because mixed botanical "
    "datasets frequently contain herbarium scans and botanical illustrations whose visual "
    "domain differs sharply from live photographs, an automated image-type screening step "
    "(based on colour statistics) was applied and its output manually curated, so that only "
    "live photographs entered training and evaluation; including herbarium specimens or line "
    "drawings would otherwise introduce domain-gap artefacts. The resulting class distribution "
    "is strongly long-tailed: the two largest genera (*Bulbophyllum*, *Dendrobium*) together "
    "account for a large fraction of all images, while the median genus is represented by "
    "approximately 21 photographs."
)

doc.add_heading("2.2 Species-stratified frozen partition", level=2)
body(
    "All images were partitioned once into training, validation, and test sets under "
    "**species-level stratification**—every photograph of a given species was assigned to a "
    "single split (seed 42)—so that no species appears in more than one partition. This is "
    "stricter than image-level splitting and prevents the model from recognising specific "
    "specimens rather than genus-level morphology. The frozen partition comprises 11,677 "
    "training, 2,629 validation, and 2,395 test photographs. Because rare genera may contribute "
    "no species to the test split, genus-level test metrics are computed over the 58 genera "
    "represented in the test partition. The identical partition was used for every model, so "
    "that differences between systems are not confounded by partition variance."
)

doc.add_heading("2.3 Stage 1 — genus classifier and backbone comparison", level=2)
body(
    "The genus classifier was selected empirically by comparing four pretrained backbones "
    "(Table 1): two Vision Transformers—DINOv2 ViT-L/14 (general self-supervised pretraining; "
    "Oquab et al., 2024) and BioCLIP 2 ViT-L/14 (hierarchical biological vision–language "
    "contrastive pretraining; Gu et al., 2025)—and two convolutional networks—ConvNeXt "
    "V2-Large (fully-convolutional masked-autoencoder pretraining; Woo et al., 2023) and "
    "EfficientNetV2-Large (supervised ImageNet-21k pretraining; Tan & Le, 2021). Each backbone "
    "was fitted with a single linear classification head over the 120 genera, applied to its "
    "pooled feature representation."
)
table(
    ["Backbone", "Family", "Pretraining", "Input (px)", "Embed. dim"],
    [["DINOv2 ViT-L/14", "ViT", "Self-supervised (natural images)", "448", "1024"],
     ["BioCLIP 2 ViT-L/14", "ViT", "Contrastive vision–language (biological)", "224", "768"],
     ["ConvNeXt V2-L", "CNN", "Masked autoencoder (ImageNet)", "384", "1536"],
     ["EfficientNetV2-L", "CNN", "Supervised (ImageNet-21k)", "448", "1280"]],
    "Table 1. The four pretrained backbones compared under an identical fine-tuning protocol."
)
body(
    "An **identical training protocol** was applied to all four backbones so the comparison is "
    "not confounded by training-regime differences. Each model was fine-tuned in two phases—a "
    "short warm-up (3 epochs) in which only the classification head is adapted, followed by "
    "end-to-end fine-tuning of the full network (up to 22 epochs, early-stopped with patience 8)—"
    "using cross-entropy loss with label smoothing (ε = 0.1) and no class-balanced resampling. "
    "Optimisation used AdamW (weight decay 0.05) with discriminative learning rates (head "
    "1×10⁻³, backbone 2×10⁻⁵) under a warm-up-then-decay schedule, and an exponential moving "
    "average of the weights (decay 0.9998) was tracked throughout. Crucially, "
    "the deployed checkpoint for each model was selected by **best validation global top-1 "
    "accuracy** rather than by validation macro-F1, which we found discarded better checkpoints. "
    "Each backbone was fine-tuned and evaluated at its native pretraining resolution (Table 1), "
    "the setting at which its pretrained weights are expected to transfer best. Models were "
    "implemented in PyTorch with the timm library and trained in the cloud (Google Colab) on a "
    "single NVIDIA L4 GPU (24 GB), using automatic mixed-precision (fp16 autocast) and "
    "multi-worker data loading; the batch size was fixed per model so that the optimisation "
    "trajectory was identical regardless of the run. Reference embeddings for Stage 2 retrieval "
    "were extracted on the same hardware under fp16 autocast."
)

doc.add_heading("2.4 Stage 2 — FAISS species retrieval", level=2)
body(
    "Because most species have too few images for direct classification, species-level "
    "suggestion is framed as visual retrieval. For the deployed backbone, the 1024-dimensional "
    "pooled feature vector (obtained from the pre-classifier representation) is extracted for "
    "every reference photograph, L2-normalised so that inner product equals cosine similarity, "
    "and indexed with FAISS using an exact inner-product index (Johnson et al., 2021). At query "
    "time, Stage 1 predicts the genus, the index is restricted to that genus, and the most "
    "similar reference photographs—together with their species labels—are returned for the "
    "user to make a final visual determination. Multiple photographs of the same specimen can be "
    "fused by confidence-weighted averaging of their embeddings before retrieval."
)

doc.add_heading("2.5 Evaluation metrics and uncertainty", level=2)
body(
    "For genus classification we report global (micro-averaged) top-1 and top-5 accuracy and, "
    "as our primary metric, **macro-averaged top-1 accuracy**, which weights every genus equally "
    "and is therefore sensitive to performance on the long tail. Probabilistic calibration is "
    "quantified by the Expected Calibration Error (ECE; Guo et al., 2017), reported both as "
    "trained and after **post-hoc temperature scaling**—a single scalar temperature T fitted on "
    "the validation split by minimising the negative log-likelihood (Guo et al., 2017); because T "
    "divides all logits equally it cannot change the ranking, so top-1 and macro accuracy are "
    "invariant and only calibration is affected. Retrieval cannot be "
    "evaluated on the species-stratified partition—every test species has, by construction, no "
    "reference photographs—so it is assessed on a separate photo-level hold-out: for every "
    "species with at least two photographs, one photograph (fixed seed) is withheld as a query "
    "and the remainder form the reference index, and we report species- and genus-level "
    "Recall@k. Every reported metric is accompanied by a 95% confidence interval obtained by "
    "non-parametric bootstrap of the test set (1,000 resamples; Efron & Tibshirani, 1993)."
)

# ── 3. RESULTS ────────────────────────────────────────────────────────────────────
doc.add_heading("3. Results", level=1)

doc.add_heading("3.1 Backbone comparison: Vision Transformers versus convolutional networks", level=2)
body(
    "On the frozen test partition (2,395 photographs; 58 genera), the self-supervised Vision "
    "Transformer DINOv2 attained the best performance on every classification metric (Table 2, "
    "Figure 2), with macro top-1 accuracy of **66.9%** (95% CI 63.7–70.6) and global top-1 "
    "of 88.9%. It was followed by the second Vision Transformer, BioCLIP 2 (macro top-1 59.8%, "
    "CI 56.2–63.1), and then the two convolutional networks, ConvNeXt V2-L (52.8%, CI "
    "49.8–56.7) and EfficientNetV2-L (47.8%, CI 44.6–51.5). The architecture families "
    "separate cleanly: **both Vision Transformers outranked both convolutional networks**, and "
    "DINOv2's confidence interval is disjoint from those of both convolutional models. The same "
    "ordering held for global top-1, macro top-5, and macro top-1 restricted to genera with at "
    "least five test photographs (DINOv2 69.6%, BioCLIP 2 60.4%, ConvNeXt V2-L 58.0%, "
    "EfficientNetV2-L 52.1%)."
)
table(
    ["Backbone", "Family", "Global T1", "Macro T1", "Macro T1 (≥5)", "Macro T5",
     "ECE (raw)", "ECE (post-TS)", "Species R@5"],
    [["DINOv2 ViT-L/14", "ViT", "88.9", "66.9", "69.6", "83.6", "0.178", "0.029", "86.6"],
     ["BioCLIP 2 ViT-L/14", "ViT", "84.4", "59.8", "60.4", "79.9", "0.205", "0.030", "86.3"],
     ["ConvNeXt V2-L", "CNN", "83.2", "52.8", "58.0", "74.4", "0.146", "0.031", "85.0"],
     ["EfficientNetV2-L", "CNN", "80.5", "47.8", "52.1", "72.2", "0.142", "0.028", "84.2"]],
    "Table 2. Backbone comparison on the frozen test partition (n = 2,395; 58 genera). Values "
    "are percentages except ECE. Macro top-1 is the primary metric; ECE (raw) is the as-trained "
    "Expected Calibration Error and ECE (post-TS) is after a single temperature scalar fitted on "
    "the validation split (Guo et al., 2017), which leaves top-1/accuracy unchanged. Best per "
    "column in bold in the text.", fontsize=8
)
figure("fig_macro_ci.png",
       "**Figure 2.** Genus-classification macro top-1 accuracy with 95% bootstrap confidence "
       "intervals (left) and Expected Calibration Error (right) for the four backbones, coloured "
       "by architecture family. Both Vision Transformers outrank both convolutional networks on "
       "accuracy, whereas the convolutional networks have lower as-trained calibration error; "
       "all four backbones reach ECE ≈ 0.03 after post-hoc temperature scaling (Figure 3, Table 2).")
body(
    "Two results stand out. First, the general-purpose self-supervised Transformer (DINOv2) "
    "**outperformed the domain-matched biological model (BioCLIP 2) by 7.1 points** of macro "
    "top-1, indicating that for this fine-grained genus task a strong self-supervised "
    "representation transfers better than vision–language contrastive pretraining on a broad "
    "biological corpus. Second, in their as-trained outputs calibration ran counter to accuracy: "
    "the convolutional networks were the best calibrated (ECE 0.142–0.146) and the Vision "
    "Transformers the worst (DINOv2 0.178, BioCLIP 2 0.205). This miscalibration was a systematic "
    "**under-confidence**—across the confidence range the empirical accuracy exceeded the stated "
    "confidence, so the reliability curve lies above the diagonal (Figure 3)—as expected when a "
    "120-way softmax distributes probability mass over many genera. A single post-hoc temperature "
    "fitted on the validation split (T < 1, which sharpens the distribution; Guo et al., 2017) "
    "corrected this for every backbone, reducing ECE to ≈0.03 (Table 2, ECE post-TS) and erasing "
    "the Vision Transformers' calibration disadvantage; because the temperature divides all logits "
    "equally it is a monotonic rescaling, so top-1 and macro accuracy are unchanged. With macro "
    "top-1 as the primary criterion, DINOv2 was selected as the deployed backbone."
)
figure("fig_reliability.png",
       "**Figure 3.** Reliability diagram for the deployed DINOv2 classifier on the test "
       "partition (15 equal-width confidence bins). The as-trained curve (red) lies above the "
       "diagonal across the confidence range—empirical accuracy exceeds the stated confidence—"
       "showing the classifier is systematically *under*-confident, a consequence of the 120-way "
       "softmax spreading probability mass across many genera. A single validation-fitted "
       "temperature (T = 0.73) sharpens the predictions onto the diagonal, lowering ECE from "
       "0.178 to 0.029 without changing any prediction (top-1 accuracy is identical).", width=4.3)

doc.add_heading("3.2 Best model in detail — DINOv2 ViT-L/14", level=2)
body(
    "Performance of the deployed model was strongly support-dependent (Figure 4). The abundant "
    "genera were near-saturated (*Bulbophyllum*, n = 808: 98.9%; *Dendrobium*, n = 622: 97.9%), "
    "while the long tail of small genera drove the gap between global (88.9%) and macro (66.9%) "
    "accuracy. This pattern—high micro accuracy with a substantial macro penalty—is the "
    "expected signature of a severely long-tailed task and motivates reporting macro accuracy as "
    "the primary metric."
)
figure("fig_dinov2_per_genus.png",
       "**Figure 4.** Per-genus top-1 accuracy of DINOv2 on the test partition, sorted by "
       "accuracy, with bars coloured by test support (number of images, log scale). Accuracy "
       "declines sharply for low-support genera in the long tail.")

doc.add_heading("3.3 Error structure and “magnet” classes", level=2)
body(
    "The confusion structure was shared across all four backbones and dominated by two abundant "
    "genera acting as attractors (Figure 5). Summed over the four models, misclassifications "
    "were absorbed overwhelmingly by *Dendrobium* (355 test images) and *Bulbophyllum* (186), "
    "the two largest classes. Several low-support genera collapsed to 0% top-1 in every model "
    "(*Aglossorrhyncha*, *Cylindrolobus*), being predicted as one of these magnets, and "
    "*Pinalia* → *Dendrobium* was the single largest confused pair (about 69% of *Pinalia* "
    "test images). DINOv2's advantage was concentrated precisely on the hardest genera (Table 3): "
    "it alone resisted the *Bulbophyllum* attractor for *Paphiopedilum* (81.0% top-1 versus "
    "≤ 9.5% for the other three backbones, which lost 76–95% of *Paphiopedilum* images to "
    "*Bulbophyllum*)."
)
figure("fig_dinov2_confusion.png",
       "**Figure 5.** Row-normalised confusion matrix for DINOv2 on the test partition. The "
       "strong diagonal indicates correct classification; vertical bands at *Dendrobium* and "
       "*Bulbophyllum* reveal their role as error “magnets” for low-support genera.",
       width=5.6)
table(
    ["Genus", "n", "Mean", "DINOv2", "BioCLIP 2", "ConvNeXt V2-L", "EffNetV2-L"],
    [["Aglossorrhyncha", "17", "0.0", "0.0", "0.0", "0.0", "0.0"],
     ["Cylindrolobus", "11", "0.0", "0.0", "0.0", "0.0", "0.0"],
     ["Anoectochilus", "19", "1.3", "5.3", "0.0", "0.0", "0.0"],
     ["Trichoglottis", "16", "1.6", "6.2", "0.0", "0.0", "0.0"],
     ["Poaephyllum", "13", "3.9", "0.0", "15.4", "0.0", "0.0"],
     ["Trachoma", "29", "10.3", "20.7", "10.3", "6.9", "3.4"],
     ["Pinalia", "36", "18.1", "27.8", "19.4", "16.7", "8.3"],
     ["Vrydagzynea", "10", "20.0", "50.0", "20.0", "10.0", "0.0"],
     ["Paphiopedilum", "21", "22.6", "81.0", "0.0", "9.5", "0.0"],
     ["Cleisostoma", "26", "33.6", "34.6", "11.5", "57.7", "30.8"]],
    "Table 3. The ten hardest genera (lowest mean top-1 across models; test support n ≥ 10). "
    "Values are top-1 accuracy (%).", fontsize=8
)

doc.add_heading("3.4 Species retrieval and embedding-space structure", level=2)
body(
    "Using DINOv2 embeddings, the retrieval stage attained species Recall@5 of 86.6% (95% CI "
    "84.7–88.4) and genus Recall@5 of 98.7% on the photo-level hold-out; the species figure "
    "narrowly exceeded that of every other backbone, so a single backbone wins both stages. A "
    "two-dimensional UMAP projection of the DINOv2 reference embeddings reveals that the "
    "representation is strongly **genus-structured without any retrieval-stage supervision**: the "
    "two most abundant genera occupy large, well-separated regions and smaller genera form "
    "compact, distinct islands (this geometry is displayed in Figure 7a for the "
    "leave-12-genera-out model of Section 3.5, whose embedding space is essentially identical to "
    "the deployed model's). This geometric separation is what makes nearest-neighbour retrieval "
    "viable for species suggestion in the absence of a species-level classifier; the diffuseness "
    "of the *Dendrobium* manifold mirrors its exceptional morphological breadth and its role as "
    "the dominant error sink in Stage 1. Strikingly, the embedding also resolves structure below "
    "the genus level: the two visually separate *Dendrobium* sub-clusters in Figure 7a correspond "
    "to a major infrageneric division—a compact island of the soft-stemmed, montane, "
    "bird-pollinated section *Calyptrochilus* (94% of all section-*Calyptrochilus* photographs, "
    "together with allied montane species such as *D. cuthbertsonii*) set apart from the main mass "
    "of larger-plant sections (*Grastidium*, *Latouria*, *Spatulata*). That section-level "
    "morphology emerges without any taxonomic supervision reinforces why embedding retrieval is "
    "effective for within-genus species suggestion; section assignments follow the "
    "orchidsnewguinea.com database and remain subject to expert revision."
)

doc.add_heading("3.5 Open-set detection of novel genera", level=2)
body(
    "The deployed model recognises 120 genera, but New Guinea harbours more, so a query may "
    "belong to a genus the system has never seen. We therefore tested whether such inputs can be "
    "flagged rather than silently misclassified, scoring each query by the cosine distance from "
    "its embedding to the nearest reference embedding—a larger distance indicating a more "
    "probable novel genus (Hendrycks & Gimpel, 2017; Vaze et al., 2022). We assessed this in two "
    "complementary ways: an optimistic *fixed-model* bound and a stricter *retraining* test."
)
body(
    "In the first, with the deployed DINOv2 model held fixed, we performed leave-one-genus-out "
    "over every genus with at least ten photographs (88 genera): each genus in turn was removed "
    "from the reference bank and its photographs scored as unknowns. Genera proved strongly "
    "separable in the retrieval embedding (mean AUROC **0.963**, median 0.977, range 0.736–0.999; "
    "99% of genera above 0.80; Figure 6). This is an *optimistic* bound, however, because the "
    "backbone had been trained on all 120 genera and may therefore carry latent structure for the "
    "“withheld” genus."
)
body(
    "To obtain a deployment-realistic estimate we ran a stricter **leave-K-genera-out retraining** "
    "test in which the withheld genera are genuinely unseen during training. Twelve genera spanning "
    "the full support range (16–274 photographs; deliberately including the morphologically "
    "distinctive slipper orchid *Paphiopedilum*) were removed from the training and validation sets, "
    "and DINOv2 was re-trained from the same initialisation under the identical protocol on the "
    "remaining 108 genera. The retrained backbone was essentially unchanged in closed-set quality "
    "(validation global top-1 85.1%), confirming that withholding rare genera did not degrade the "
    "representation. All 16,701 photographs—including the twelve unseen genera as queries—were then "
    "embedded with this hold-out model and each scored by cosine distance to its nearest *known* "
    "reference. The genuinely-unseen genera remained highly separable: mean per-genus AUROC "
    "**0.958** (median 0.966, range 0.894–0.997; pooled 0.961), with *Paphiopedilum* the hardest "
    "(0.894) and *Apostasia* the easiest (0.997)."
)
body(
    "A genus-by-genus comparison against the optimistic fixed-model bound (Table 4) makes this "
    "robustness concrete. Averaged over the same twelve genera, the strict estimate (0.958) was, if "
    "anything, marginally **higher** than the fixed-model bound (0.955; mean change +0.003, mean "
    "absolute change 0.018), and for five of the twelve genera open-set detection actually improved "
    "when the genus was genuinely unseen; only *Paphiopedilum* showed a substantial decline "
    "(−0.049). That the strict, retrained estimate **matches**—rather than collapses below—the "
    "fixed-model bound (pooled 0.961 vs 0.963) indicates that the separability reflects genuine "
    "representational structure of the DINOv2 embedding rather than memorised training genera. Two "
    "caveats temper a literal reading of the per-genus differences. First, the two protocols are "
    "**not strictly identical**: the fixed-model test scores a single genus as unknown against the "
    "remaining 119 known genera, whereas the strict test scores all twelve withheld genera as "
    "unknown against only 108 known genera, so the reference set and the in-distribution pool differ "
    "and the comparison is practical rather than perfectly controlled. Second, per-genus AUROC for "
    "the rarest genera carries appreciable sampling uncertainty. The near-equality nonetheless holds "
    "across both the pooled estimate and the paired per-genus values. Figure 7a shows the held-out "
    "*Paphiopedilum* photographs forming a compact, peripheral island in the embedding space, and "
    "Figure 7b that they sit at markedly larger distances from the known set than in-distribution "
    "photographs. As a concrete deployment operating point, we fix the threshold τ at the 95th "
    "percentile of the known-genus nearest-neighbour distances (τ = 0.29 in cosine distance, a 5% "
    "false-positive rate on known genera), at which 70% of genuinely novel-genus photographs are "
    "flagged; relaxing the tolerance to a 10% false-positive rate (τ = 0.21) raises this to 94%. "
    "In practice this lets the system abstain on probable novel genera and route them to expert "
    "review through the Lucid key, converting a silent error into an explicit “unknown—needs "
    "verification” outcome."
)
figure("fig_openset_auroc_hist.png",
       "**Figure 6.** Open-set detection of novel genera by nearest-reference embedding distance "
       "(leave-one-genus-out, fixed DINOv2 model). Distribution of per-genus AUROC across 88 "
       "genera (n ≥ 10); mean 0.963. Most genera are highly separable from the known set.",
       width=5.6)
table(
    ["Held-out genus", "n", "Fixed-model AUROC", "Strict AUROC", "Δ (strict − fixed)"],
    [["Apostasia", "16", "0.998", "0.997", "−0.002"],
     ["Pterostylis", "47", "0.971", "0.980", "+0.009"],
     ["Chrysoglossum", "21", "0.981", "0.975", "−0.006"],
     ["Aglossorrhyncha", "26", "0.974", "0.970", "−0.004"],
     ["Oberonia", "274", "0.987", "0.970", "−0.018"],
     ["Robiquetia", "172", "0.952", "0.966", "+0.014"],
     ["Thrixspermum", "75", "0.979", "0.966", "−0.014"],
     ["Crepidium", "105", "0.963", "0.962", "−0.001"],
     ["Spiranthes", "21", "0.912", "0.955", "+0.044"],
     ["Dryadorchis", "37", "0.927", "0.938", "+0.011"],
     ["Vanda", "44", "0.878", "0.927", "+0.049"],
     ["Paphiopedilum", "51", "0.944", "0.894", "−0.049"],
     ["Mean (12 genera)", "—", "0.955", "0.958", "+0.003"]],
    "Table 4. Paired open-set AUROC for the twelve held-out genera under both protocols: the "
    "optimistic fixed-model bound (leave-one-genus-out; DINOv2 trained on all 120 genera, the genus "
    "removed only from the reference bank) versus the strict retraining test (DINOv2 retrained on "
    "the remaining 108 genera, the genus genuinely unseen during training). Δ is strict minus fixed; "
    "positive values mean the genus was detected as novel at least as well when genuinely unseen. "
    "The two protocols are not strictly identical (fixed: one unknown genus vs 119 known; strict: "
    "twelve unknown vs 108 known), so the comparison is practical rather than perfectly controlled. "
    "Mean over the twelve genera: fixed 0.955, strict 0.958.", fontsize=8
)
figure("fig_openset_umap_combined.png",
       "**Figure 7.** Embedding-space structure and open-set separation under the strict "
       "leave-12-genera-out retraining test. (a) UMAP projection (cosine metric) of the 16,701 "
       "DINOv2 embeddings from the hold-out model, with the twelve most photographed known genera "
       "coloured, remaining genera in grey, and the unseen exemplar Paphiopedilum overlaid as red "
       "stars; the unseen genus forms a compact peripheral island. (b) Cosine distance to the "
       "nearest known reference embedding for in-distribution (known) photographs versus the unseen "
       "Paphiopedilum, which separates from the known set with AUROC 0.89.",
       width=6.6)

# ── 4. DISCUSSION ─────────────────────────────────────────────────────────────────
doc.add_heading("4. Discussion", level=1)
body(
    "Under identical, leakage-free conditions, **Vision Transformers clearly outperformed "
    "convolutional networks** for fine-grained orchid genus recognition: both Transformers "
    "ranked above both convolutional models on the primary metric, with confidence intervals "
    "that do not overlap between the best Transformer and either convolutional network. This is "
    "consistent with the view that the global receptive field of self-attention (Dosovitskiy et "
    "al., 2021) is well-suited to capturing the spatially distributed floral micro-characters "
    "that distinguish orchid genera, although we note that input resolution is partially "
    "entangled with architecture family (Section 4, Limitations)."
)
body(
    "More surprising is that **general-purpose self-supervised pretraining outperformed "
    "domain-matched biological pretraining**. DINOv2, trained on curated natural images without "
    "labels (Oquab et al., 2024), exceeded BioCLIP 2 (Gu et al., 2025) by 7.1 points of macro "
    "top-1, despite the latter's biological vision–language training. A plausible explanation "
    "is that contrastive vision–language objectives optimise for cross-modal alignment and "
    "broad taxonomic coverage rather than for the within-family visual discrimination required "
    "here, whereas DINOv2's dense self-supervised features retain finer local detail. This "
    "echoes the broader finding that strong self-supervised features transfer remarkably well "
    "across domains, and cautions against assuming that domain-matched foundation models are "
    "always preferable for narrow downstream tasks."
)
body(
    "The error analysis exposes a **long-tailed “magnet-class” failure mode** of practical "
    "importance: misclassifications concentrate on the two most abundant genera, and several "
    "rare genera are never recovered. This is the expected behaviour of cross-entropy training "
    "under extreme class imbalance (Cui et al., 2019) and suggests that targeted "
    "interventions—class-balanced objectives, additional reference imagery for the most "
    "absorbed genera, or hierarchical decision rules—could yield further gains beyond the "
    "choice of backbone. Encouragingly, the best backbone is also the most robust on the hardest "
    "genera, indicating that representation quality, not merely class frequency, governs "
    "tail performance."
)
body(
    "The two-stage design directly addresses data scarcity. Rather than attempting infeasible "
    "1,350-way species classification, the system predicts the more learnable genus and defers "
    "species determination to embedding retrieval and, ultimately, to the user. The strong genus "
    "structure of the DINOv2 embedding space (Figure 7a) and the high genus-level Recall@5 "
    "(98.7%) make this division of labour effective. In deployment, genus predictions are linked "
    "to a Lucid multi-access identification key, allowing users to confirm a determination with "
    "morphological characters—combining the throughput of machine vision with the rigour of "
    "expert keys. Robustness to the inevitable arrival of unrecognised genera is provided by the "
    "open-set distance score (Section 3.5): because that score remained discriminative even for "
    "genera entirely unseen during training (strict mean AUROC 0.958, essentially matching the "
    "fixed-model bound), the system can abstain on probable novel genera and defer them to expert "
    "review rather than misclassify them silently."
)
body(
    "**Limitations.** First, each backbone was evaluated at its native pretraining resolution "
    "(224–448 px), so input resolution is partially confounded with architecture family; a "
    "resolution-matched ablation would further isolate the architecture effect. Second, in their "
    "as-trained outputs the Vision Transformers were less well calibrated (more under-confident) "
    "than the convolutional networks; post-hoc temperature scaling reduced every backbone's ECE "
    "to ≤0.03 (Section 3.1), so this is readily corrected, but the temperature should be re-fitted "
    "if the model or input distribution changes. Third, retrieval quality is bounded by the "
    "coverage and label accuracy of the "
    "reference set; species absent from the reference index cannot be retrieved. Finally, the "
    "evaluation covers the 58 genera represented in a single frozen test partition; broader "
    "geographic and taxonomic validation is desirable."
)

# ── 5. CONCLUSION ─────────────────────────────────────────────────────────────────
doc.add_heading("5. Conclusion", level=1)
body(
    "For fine-grained orchid identification in a species-rich, data-poor flora, a "
    "self-supervised Vision-Transformer backbone (DINOv2) combined with embedding-based species "
    "retrieval is an effective and deployable strategy. In a controlled benchmark, Vision "
    "Transformers outperformed convolutional networks, and general-purpose self-supervised "
    "pretraining outperformed domain-matched biological pretraining. The resulting two-stage "
    "system—released as an open web application linked to a multi-access identification key—offers a "
    "practical template for automated identification in other hyperdiverse, under-documented "
    "taxa."
)

# ── BACK MATTER ───────────────────────────────────────────────────────────────────
doc.add_heading("Data and Code Availability", level=1)
body(
    "Training and evaluation code, the deployed model weights, and the interactive application "
    "are openly available (web application on Hugging Face Spaces; code repository linked "
    "therein). The curated photographic dataset is described in detail above; reference imagery "
    "originates from orchidsnewguinea.com and contributed field collections and is available "
    "subject to the originators' terms."
)

doc.add_heading("Acknowledgements", level=1)
body(
    "We thank the contributors and curators of orchidsnewguinea.com and the field "
    "photographers whose images made this work possible, and the maintainers of the Lucid "
    "multi-access key for New Guinea orchids. This work was supported by the Australian Orchid "
    "Foundation."
)

doc.add_heading("References", level=1)
refs = [
    "Affouard, A., Goëau, H., Bonnet, P., Lombardo, J.-C., & Joly, A. (2017). Pl@ntNet app in the era of deep learning. In ICLR 2017 Workshop Track.",
    "Apriyanti, D. H., Spreeuwers, L. J., & Lucas, P. J. F. (2023). Deep neural networks for explainable feature extraction in orchid identification. Applied Intelligence, 53, 26270–26285. https://doi.org/10.1007/s10489-023-04880-2",
    "Arwatchananukul, S., Kirimasthong, K., & Aunsri, N. (2020). A new Paphiopedilum orchid database and its recognition using convolutional neural network. Wireless Personal Communications, 115, 3275–3289. https://doi.org/10.1007/s11277-020-07463-3",
    "Cámara-Leret, R., Frodin, D. G., Adema, F., et al. (2020). New Guinea has the world's richest island flora. Nature, 584, 579–583. https://doi.org/10.1038/s41586-020-2549-5",
    "Cui, Y., Jia, M., Lin, T.-Y., Song, Y., & Belongie, S. (2019). Class-balanced loss based on effective number of samples. In CVPR (pp. 9268–9277).",
    "de Vogel, E. F., Vermeulen, J. J., & Schuiteman, A. (2014). Flora Malesiana: Orchids of New Guinea. Naturalis Biodiversity Center, Leiden.",
    "Dosovitskiy, A., Beyer, L., Kolesnikov, A., et al. (2021). An image is worth 16×16 words: Transformers for image recognition at scale. In ICLR.",
    "Efron, B., & Tibshirani, R. J. (1993). An Introduction to the Bootstrap. Chapman & Hall.",
    "Gu, J., Stevens, S., Campolongo, E. G., et al. (2025). BioCLIP 2: Emergent properties from scaling hierarchical contrastive learning. arXiv:2505.23883.",
    "Guo, C., Pleiss, G., Sun, Y., & Weinberger, K. Q. (2017). On calibration of modern neural networks. In ICML (pp. 1321–1330).",
    "Hendrycks, D., & Gimpel, K. (2017). A baseline for detecting misclassified and out-of-distribution examples in neural networks. In ICLR.",
    "Johnson, J., Douze, M., & Jégou, H. (2021). Billion-scale similarity search with GPUs. IEEE Transactions on Big Data, 7(3), 535–547.",
    "Oquab, M., Darcet, T., Moutakanni, T., et al. (2024). DINOv2: Learning robust visual features without supervision. Transactions on Machine Learning Research.",
    "Schuiteman, A. (1995). Key to the genera of Orchidaceae of New Guinea. Flora Malesiana Bulletin, 11(6), 401–424.",
    "Stevens, S., Wu, J., Thompson, M. J., et al. (2024). BioCLIP: A vision foundation model for the tree of life. In CVPR (pp. 19412–19424).",
    "Tan, M., & Le, Q. (2021). EfficientNetV2: Smaller models and faster training. In ICML (pp. 10096–10106).",
    "Van Horn, G., Mac Aodha, O., Song, Y., et al. (2018). The iNaturalist species classification and detection dataset. In CVPR (pp. 8769–8778).",
    "Vaze, S., Han, K., Vedaldi, A., & Zisserman, A. (2022). Open-set recognition: A good closed-set classifier is all you need? In ICLR.",
    "Vitt, P., Taylor, A., Rakosy, D., et al. (2023). Global conservation prioritization for the Orchidaceae. Scientific Reports, 13, 6718. https://doi.org/10.1038/s41598-023-30177-y",
    "Woo, S., Debnath, S., Hu, R., et al. (2023). ConvNeXt V2: Co-designing and scaling ConvNets with masked autoencoders. In CVPR (pp. 16133–16142).",
]
for r in refs:
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(r)
    run.font.size = Pt(10)

add_page_numbers()
OUT.parent.mkdir(parents=True, exist_ok=True)
target = OUT
try:
    doc.save(str(target))
except PermissionError:
    target = OUT.with_name(OUT.stem + "_NEW.docx")
    doc.save(str(target))
    print("(original was locked/open — saved to a new file instead)")
print(f"Saved manuscript → {target}")
print(f"Sections: headings populate the Word Navigation pane. Figures embedded: 7. Tables: 4. References: {len(refs)}.")
